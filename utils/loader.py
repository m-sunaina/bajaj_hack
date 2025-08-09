import os
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

def load_and_chunk(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100
    )

    chunks = []

    if ext == ".pdf":
        # Lazy import to speed up startup
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        chunks = splitter.split_documents(docs)

        for chunk in chunks:
            page_number = chunk.metadata.get("page", "?")
            try:
                page_number = int(page_number)
            except:
                page_number = None

            chunk.metadata = {
                "source": filename,
                "page": page_number
            }

    elif ext == ".docx":
        # Lazy import to speed up startup
        from docx import Document as DocxDocument
        docx = DocxDocument(file_path)
        raw_text = "\n".join([para.text for para in docx.paragraphs if para.text.strip()])
        text_chunks = splitter.split_text(raw_text)

        chunks = [
            Document(
                page_content=chunk,
                metadata={"source": filename, "page": i + 1}
            )
            for i, chunk in enumerate(text_chunks)
        ]

    else:
        raise ValueError("Unsupported file format. Only .pdf and .docx are supported.")

    print(f"📄 Loaded '{filename}' into {len(chunks)} chunks.")
    return chunks
